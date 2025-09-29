from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

def generate_document(Nomer1, Nomer2, Gos_TS, FIO, seal_path, output_file="document.docx"):
    """
    Генерирует маршрутный лист с титульником, датой слева и печатью справа внизу страницы.
    """
    doc = Document()

    # Титульная страница
    title = doc.add_heading(f'Маршрутный Лист {Nomer1}', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(f'Открыт по Путевому Листу {Nomer2}').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # пустая строка

    # Пустое пространство перед нижней частью
    for _ in range(2):
        doc.add_paragraph()

    # Основное содержание
    doc.add_paragraph(f'Техническое состояние ТС: {Gos_TS} в удовлетворительном состоянии.')
    doc.add_paragraph(f'ЭМО водителем {FIO} пройден успешно.')

    # Пустое пространство перед нижней частью
    for _ in range(7):
        doc.add_paragraph()

    # Нижняя часть с таблицей: дата слева, печать справа
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3)
    table.columns[1].width = Inches(3)

    # Левый столбец: дата и подпись ниже даты
    today = datetime.today().strftime('%d.%m.%Y')
    left_cell = table.cell(0, 0)
    left_cell.text = f"Дата: {today}\n\nПодпись: ____________"
    for paragraph in left_cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Правый столбец: печать
    right_cell = table.cell(0, 1)
    paragraph = right_cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    try:
        paragraph.add_run().add_picture(seal_path, width=Inches(3))  # увеличиваем печать
    except Exception as e:
        right_cell.text = f"Ошибка при вставке печати: {e}"

    doc.save(output_file)
    print(f"Документ '{output_file}' успешно создан.")
